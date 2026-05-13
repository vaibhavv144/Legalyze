"""
One-off generator for Legal_AI complete technical DOCX. Run from repo root:
  pip install python-docx matplotlib   # matplotlib optional (embeds diagrams)
  python _build_documentation_docx.py

Output: Legal_AI_Complete_Technical_Documentation.docx
"""
from __future__ import annotations

import os
import tempfile

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_together = True
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "EDEDED")
    p._p.get_or_add_pPr().append(shd)


def add_h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_paragraph(text, style=f"Heading {level}")


def add_p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = val


def page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _render_diagram(path: str, title: str, boxes: list[tuple[float, float, float, float, str]], arrows: list[tuple[float, float, float, float]]) -> bool:
    """Draw a simple block diagram with matplotlib (non-interactive). Returns False if matplotlib missing."""
    try:
        import matplotlib

        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
    except ImportError:
        return False

    fig, ax = plt.subplots(figsize=(11, 6.5))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 8)
    ax.axis("off")
    ax.set_title(title, fontsize=13, fontweight="bold", pad=12)

    for x, y, w, h, text in boxes:
        patch = FancyBboxPatch(
            (x, y),
            w,
            h,
            boxstyle="round,pad=0.02,rounding_size=0.15",
            linewidth=1.4,
            edgecolor="#1e40af",
            facecolor="#dbeafe",
        )
        ax.add_patch(patch)
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=9, wrap=True)

    for x1, y1, x2, y2 in arrows:
        arr = FancyArrowPatch(
            (x1, y1),
            (x2, y2),
            arrowstyle="-|>",
            mutation_scale=14,
            linewidth=1.2,
            color="#334155",
        )
        ax.add_patch(arr)

    plt.tight_layout()
    plt.savefig(path, dpi=140, bbox_inches="tight", facecolor="white")
    plt.close()
    return True


def add_embedded_diagrams(doc: Document, tmpdir: str) -> None:
    add_h(doc, "Visual diagrams (system, RAG, authentication)", 1)
    add_p(
        doc,
        "These figures are rendered automatically when matplotlib is installed. "
        "If you only install python-docx, the textual ASCII diagrams elsewhere still explain the flows.",
    )

    rag_png = os.path.join(tmpdir, "rag_pipeline.png")
    ok_rag = _render_diagram(
        rag_png,
        "Figure 1 — RAG chat path (this codebase)",
        [
            (0.6, 4.8, 2.2, 1.0, "User message\n(chat route)"),
            (3.2, 4.8, 2.1, 1.0, "Embed query\n(text-embedding-004)"),
            (5.7, 4.8, 2.4, 1.0, "Load user chunks\n(Mongo ≤3000)"),
            (8.6, 4.8, 2.2, 1.0, "Score chunks\n0.65·cos +\n0.35·lexical"),
            (10.9, 4.8, 2.2, 1.0, "Top-k context\nbuild_context"),
            (5.6, 2.6, 2.8, 1.0, "Gemini generateContent\n(legal_chat prompt)"),
            (5.6, 0.9, 2.8, 1.0, "Assistant reply +\ncitations (UI)"),
        ],
        [
            (2.8, 5.3, 3.2, 5.3),
            (5.3, 5.3, 5.7, 5.3),
            (8.1, 5.3, 8.6, 5.3),
            (10.85, 5.3, 11.15, 4.85),
            (11.95, 4.65, 7.05, 3.45),
            (7.0, 2.6, 7.0, 1.95),
        ],
    )

    sys_png = os.path.join(tmpdir, "system_arch.png")
    ok_sys = _render_diagram(
        sys_png,
        "Figure 2 — Logical deployment (Docker / local)",
        [
            (0.5, 5.6, 2.6, 1.25, "Browser\nReact + Vite\n(or nginx static)"),
            (4.2, 5.6, 2.8, 1.25, "FastAPI\nUvicorn :8000\nREST + JWT"),
            (8.5, 5.6, 2.5, 1.25, "MongoDB 7\nCollections:\nusers, docs, chats…"),
            (4.0, 3.6, 3.4, 1.05, "Disk: uploads/\nPDF, DOCX, images"),
            (4.0, 1.9, 3.4, 1.05, "Google Gemini API\n(generate + embed)\noptional key"),
        ],
        [
            (3.15, 6.23, 4.15, 6.23),
            (7.1, 6.23, 8.45, 6.23),
            (5.6, 5.58, 5.6, 4.73),
            (5.6, 3.53, 5.6, 2.93),
        ],
    )

    auth_png = os.path.join(tmpdir, "jwt_flow.png")
    ok_auth = _render_diagram(
        auth_png,
        "Figure 3 — JWT access pattern (high level)",
        [
            (0.6, 5.8, 2.4, 1.05, "POST /auth/login\nemail + password"),
            (4.4, 5.8, 2.4, 1.05, "Verify hash\n(passlib)"),
            (8.4, 5.8, 2.9, 1.05, "Issue JWTs\n(access + refresh)\nHS256 + exp"),
            (1.9, 3.85, 2.8, 1.05, "Frontend stores\ntokens (localStorage)"),
            (6.25, 3.85, 3.9, 1.05, "Protected calls:\nAuthorization: Bearer <access>\nDepends(get_current_user_id)"),
            (4.0, 1.85, 4.35, 1.05, "Expiry → POST /auth/refresh\nwith refresh token"),
        ],
        [
            (3.05, 6.33, 4.35, 6.33),
            (7.85, 6.33, 8.35, 6.33),
            (10.85, 5.76, 3.95, 4.73),
            (4.76, 4.35, 6.21, 4.33),
            (7.08, 3.82, 6.08, 2.93),
        ],
    )

    for caption, png, ok in [
        ("Figure 1 — Retrieval-augmented chat ends to end.", rag_png, ok_rag),
        ("Figure 2 — Services and persistence.", sys_png, ok_sys),
        ("Figure 3 — Authentication and Bearer usage.", auth_png, ok_auth),
    ]:
        if ok and os.path.isfile(png):
            add_h(doc, caption, 2)
            doc.add_picture(png, width=Inches(6.35))
            add_p(doc, "")
        else:
            add_h(doc, caption + " (skipped — pip install matplotlib)", 2)


def main() -> None:
    doc = Document()
    doc.core_properties.title = "Indian Legal AI Platform — Complete Technical Documentation"
    doc.core_properties.subject = "Legal_AI project documentation"
    doc.core_properties.keywords = "FastAPI, MongoDB, RAG, Gemini, React, Legal AI"

    doc.add_paragraph("Indian Legal AI Platform", style="Title")
    doc.add_paragraph(
        "Exhaustive technical documentation generated from the repository source tree "
        "(application code only; dependencies such as node_modules and .venv are excluded)."
    )
    add_h(doc, "Table of Contents", 1)
    for item in [
        "Section 1 — Project Overview",
        "Section 2 — Technology Stack",
        "Section 3 — RAG Deep Dive",
        "Section 4 — Python Concepts Used",
        "Section 5 — Complete Code Walkthrough",
        "Section 6 — API / Endpoints",
        "Section 7 — Data Flow Diagram",
        "Section 8 — Setup & Configuration",
        "Section 9 — Master Interview Preparation",
        "Section 10 — MongoDB Collections & Fields (reference)",
        "Section 11 — Beginner Concepts (Python, HTTP, JWT, async explained)",
        "Section 12 — RAG Alternatives & Design Trade-offs",
        "Section 13 — Security, Privacy & Limitations",
        "Section 14 — Glossary",
    ]:
        doc.add_paragraph(item, style="List Number")
    add_p(
        doc,
        "Note: In Microsoft Word, you can replace this list with Insert → Table of Contents "
        "if you apply Heading styles consistently (this document does).",
    )
    page_break(doc)
    with tempfile.TemporaryDirectory(prefix="legal_ai_docx_") as tmpdir:
        add_embedded_diagrams(doc, tmpdir)
    page_break(doc)

    # ----- SECTION 1 -----
    add_h(doc, "Section 1: Project Overview", 1)
    add_h(doc, "What this project does", 2)
    add_p(
        doc,
        "This repository implements an Indian-law-focused legal assistant web application. "
        "Users register and log in with email and password. They upload contracts (PDF, DOCX, or images), "
        "and the backend extracts text, stores documents in MongoDB, and optionally chunks text for "
        "retrieval-augmented generation (RAG) using Google Gemini embeddings and cosine similarity, "
        "with lexical token overlap as a secondary signal. Chat lets users ask legal-style questions; "
        "the backend retrieves relevant chunks, builds a context string, and calls Gemini for an answer "
        "(or falls back to rule-based text when no API key is set). A separate “analysis” endpoint scans "
        "contract text with keyword and regex heuristics to produce extracted clause snippets, risk flags, "
        "and a simple summary. The frontend is a React + TypeScript single-page app styled with Tailwind "
        "CSS, offering dashboard, upload, documents with analysis views, chat, profile, and settings "
        "including global search and notifications.",
    )
    add_h(doc, "Who would use it and why", 2)
    add_p(
        doc,
        "Legal teams, founders, or procurement staff who need a lightweight MVP to organize contracts, "
        "get heuristic risk highlights, and ask questions grounded (when embeddings work) in uploaded text. "
        "It is framed as educational guidance, not a substitute for professional legal advice.",
    )
    add_h(doc, "High-level architecture (ASCII)", 2)
    add_code_block(
        doc,
        """                          +------------------+
                          |   Web Browser    |
                          |  React + Vite    |
                          +--------+---------+
                                   | HTTPS / REST (axios)
                                   v
+------------------+      +------------------+      +------------------+
|  Nginx (Docker)  |      | FastAPI (Python) |      | MongoDB 7        |
|  serves static   | ---> |  JWT auth        | ---> | users, documents |
|  frontend build  |      |  document upload |      | clauses, risks,  |
+------------------+      |  RAG + Gemini    |      | chats, chunks,   |
                          +------------------+      | notifications    |
                                   |                +------------------+
                                   v
                          Google Gemini API
                          (generateContent + embedContent)
""",
    )
    add_h(doc, "Complete folder structure (project files only)", 2)
    add_code_block(
        doc,
        """Legal_AI/
  README.md                 # Run instructions (Docker + local)
  docker-compose.yml        # mongo + backend + frontend services
  backend/
    Dockerfile
    requirements.txt
    .env.example              # Backend configuration template
    app/
      main.py                 # FastAPI app, CORS, routers, lifespan
      core/
        config.py             # Pydantic Settings (env)
        deps.py               # get_db, JWT dependency
        security.py           # hash/verify password, JWT create/decode
      db/
        mongo.py              # Motor async Mongo client
        indexes.py            # Collection indexes
      modules/
        auth/routes.py        # Register, login, refresh, /me, PATCH /users/me
        documents/routes.py   # Upload, list, get, delete, analyze
        analysis/routes.py    # Clauses, risks, summary, explain
        chat/routes.py        # Chat sessions and messages (RAG + AI)
        search/routes.py      # Regex search across entities
        notifications/routes.py
      schemas/
        common.py             # User and auth DTOs
        legal.py              # Legal domain response models
      services/
        ai_service.py         # Gemini HTTP + heuristics / fallbacks
        rag_service.py        # Chunking, embedding, hybrid retrieval
        document_service.py   # Save uploads, extract PDF/DOCX/image stub
    tests/
      test_health.py          # Health endpoint smoke test
    uploads/                  # Stored uploaded files (runtime)
  frontend/
    Dockerfile                # Multi-stage: npm build + nginx
    package.json
    vite.config.ts
    index.html
    .env.example
    src/
      main.tsx                # React root, router, react-query, toaster
      App.tsx                 # Routes and auth guard
      lib/api.ts              # Axios instance + refresh helper
      lib/types.ts            # Shared TS types
      lib/utils.ts            # cn() class merge
      hooks/use-auth.tsx      # Auth context
      pages/                  # Landing, auth, workspace pages
      components/             # UI shell, chat, upload, marketing
  shared/contracts/legal.ts   # Shared clause type unions (TS)
""",
    )
    add_p(
        doc,
        "Excluded from documentation scope: backend/.venv, frontend/node_modules, frontend/dist build artifacts "
        "(except noting they exist), and binary uploads under backend/uploads/.",
    )
    page_break(doc)

    # ----- SECTION 2 ----- (abbreviated narrative + tables per major tech)
    add_h(doc, "Section 2: Technology Stack", 1)
    add_p(
        doc,
        "Below, each major dependency is explained for readers new to Python or web development, "
        "with project-specific usage and interview-style practice.",
    )

    stacks = [
        (
            "FastAPI",
            "A modern Python web framework that builds REST APIs using standard Python type hints. "
            "It automatically validates request bodies with Pydantic and generates OpenAPI documentation.",
            "Defines HTTP routes in app.main and app.modules.*.routes, wires routers, and applies CORS.",
            "Like a restaurant counter: each endpoint is a menu item; FastAPI checks orders (requests) "
            "and plates responses.",
            [
                ("Router", "Groups related endpoints under a prefix"),
                ("Dependency injection", "Depends() injects DB and current user"),
                ("OpenAPI", "Interactive docs at /docs"),
            ],
            [
                (
                    "How does FastAPI know the shape of JSON bodies?",
                    "Through Pydantic models used as parameters or response_model.",
                ),
                (
                    "What is the difference between sync and async route functions?",
                    "Async routes can await I/O (DB, HTTP); sync routes block the worker thread.",
                ),
                (
                    "Where is CORS configured in this project?",
                    "CORSMiddleware in app/main.py allows all origins for development.",
                ),
                (
                    "How are routes organized?",
                    "Separate modules under app/modules/<area>/routes.py included in main.",
                ),
                (
                    "What generates API docs?",
                    "FastAPI builds OpenAPI schema; Swagger UI is at /docs.",
                ),
            ],
        ),
        (
            "Uvicorn",
            "An ASGI server that runs async Python web apps such as FastAPI.",
            "Docker CMD and local README use uvicorn app.main:app to serve the API.",
            "Like an engine that runs the FastAPI application process and listens on a port.",
            [
                ("ASGI", "Async Server Gateway Interface for Python web apps"),
                ("Workers", "Can scale with multiple processes in production"),
            ],
            [
                ("What is ASGI?", "A standard interface between async Python apps and servers like Uvicorn."),
                ("Why uvicorn for FastAPI?", "FastAPI is an ASGI app; uvicorn is a common production server."),
                ("How is the app object referenced?", "app.main:app means module path app.main, variable app."),
                ("What port does Docker expose?", "8000 inside the container per Dockerfile."),
            ],
        ),
        (
            "Motor + MongoDB",
            "Motor is the official async Python driver for MongoDB. MongoDB is a document database storing JSON-like documents.",
            "connect_to_mongo uses AsyncIOMotorClient; routes query collections like db.users, db.documents.",
            "MongoDB is like filing cabinets with flexible folders; Motor is the async librarian fetching papers.",
            [
                ("Collection", "Analogous to a SQL table but schema-flexible"),
                ("Index", "Speeds lookups; unique index on users.email"),
            ],
            [
                ("Why Motor instead of PyMongo directly?", "Native async/await integration with FastAPI routes."),
                ("How is the database selected?", "client[settings.mongo_db_name] in mongo.py."),
                ("What indexes are created?", "See db/indexes.py for users, documents, clauses, risks, chats, chunks."),
            ],
        ),
        (
            "Pydantic & pydantic-settings",
            "Pydantic validates Python data classes; pydantic-settings loads configuration from environment variables and .env files.",
            "Settings in core/config.py defines app_name, mongo_uri, jwt_secret, gemini keys, etc.",
            "Settings are a checklist that refuses to serve the app with inconsistent types—like forms that validate before submit.",
            [
                ("BaseSettings", "Loads fields from environment with optional .env file"),
                ("Field validation", "EmailStr, min_length on passwords in schemas"),
            ],
            [
                ("Where is .env loaded?", "SettingsConfigDict(env_file='.env') in Settings."),
                ("Why separate schemas/common.py and legal.py?", "Auth models vs legal domain models."),
            ],
        ),
        (
            "PyJWT & passlib",
            "PyJWT creates and verifies JSON Web Tokens. Passlib hashes passwords (here pbkdf2_sha256).",
            "Login returns access and refresh tokens; get_current_user_id decodes access tokens.",
            "JWT is a stamped ticket showing identity; hashing prevents storing raw passwords.",
            [
                ("Access vs refresh", "Access for API calls; refresh to obtain new access tokens"),
                ("HS256", "Symmetric signing with jwt_secret"),
            ],
            [
                ("Why not store plaintext passwords?", "If the DB leaks, attackers still lack passwords."),
                ("What claim distinguishes token types?", "Custom 'type' field: access or refresh."),
            ],
        ),
        (
            "httpx",
            "Async-capable HTTP client for Python used to call external REST APIs.",
            "AIService and RAGService use httpx.AsyncClient to call Google Gemini generateContent and embedContent.",
            "Like requests but fits async FastAPI routes without blocking.",
            [("AsyncClient", "Use with await for non-blocking I/O")],
            [
                ("How are Gemini failures handled?", "Retry once on some HTTP errors; else fallback logic."),
            ],
        ),
        (
            "python-docx & pypdf & Pillow",
            "python-docx reads Word files; pypdf reads PDF pages; Pillow is imported indirectly for image pipeline readiness.",
            "document_service.extract_text uses PdfReader and Docx paragraphs; images return a placeholder string.",
            "Like printers in reverse: turning files back into plain text the AI pipeline can read.",
            [],
            [
                ("What happens for scanned PDFs?", "Text extraction may be empty unless OCR is added; images use placeholder text."),
            ],
        ),
        (
            "React 19 + TypeScript + Vite",
            "React builds UI as components; TypeScript adds static types; Vite bundles for dev and production builds.",
            "src/App.tsx routes pages; workspace-pages.tsx contains dashboard, chat, documents, etc.",
            "React is LEGO bricks for UI; Vite is a fast factory assembling them.",
            [
                ("Hooks", "useState, useEffect, useMemo, useContext in this codebase"),
                ("react-router-dom", "BrowserRouter, Routes, Navigate for SPA navigation"),
            ],
            [
                ("Why axios interceptor?", "Attaches Bearer token from localStorage to each request."),
            ],
        ),
        (
            "Tailwind CSS 4",
            "Utility-first CSS: classes like flex, p-4, bg-slate-900 compose layouts without writing separate CSS files for each rule.",
            "index.css and components use Tailwind classes for the dark legal workspace theme.",
            "Like naming every paint chip instead of mixing paint from scratch each time.",
            [],
            [
                ("What is cn() in lib/utils.ts?", "Merges class names with tailwind-merge to resolve conflicts."),
            ],
        ),
        (
            "TanStack React Query",
            "A data-fetching library for caching and refreshing server state (used here via QueryClientProvider).",
            "QueryClient created in main.tsx; many pages still call axios directly—both can coexist.",
            "Like a smart clipboard that remembers server answers for a short time.",
            [],
            [
                ("Is every screen using useQuery?", "Not necessarily; many components call api.get in useEffect."),
            ],
        ),
        (
            "Google Gemini (REST)",
            "Gemini is Google's multimodal model family accessed via HTTP from this project (not the Python SDK).",
            "ai_service calls generateContent; rag_service calls text-embedding-004 embedContent.",
            "Two booths: one writes answers, one turns text into vectors for similarity.",
            [
                ("Model name", "settings.gemini_model default gemini-1.5-flash"),
                ("Embedding model", "text-embedding-004 in rag_service._embed_text"),
            ],
            [
                ("What happens if GEMINI_API_KEY is empty?", "LLM and embeddings return None; fallbacks and lexical-only retrieval apply."),
            ],
        ),
    ]

    for name, what, how_proj, analogy, terms, iqs in stacks:
        add_h(doc, name, 2)
        add_p(doc, "What is it? " + what)
        add_p(doc, "Why / how in this project: " + how_proj)
        add_p(doc, "Analogy: " + analogy)
        if terms:
            add_table(doc, ["Term", "Meaning"], terms)
        add_h(doc, f"Interview questions — {name}", 3)
        for q, a in iqs:
            doc.add_paragraph(q, style="List Number")
            doc.add_paragraph(a)

    page_break(doc)

    # ----- SECTION 3 RAG -----
    add_h(doc, "Section 3: RAG (Retrieval-Augmented Generation) — Deep Dive", 1)
    add_p(
        doc,
        "RAG augments a language model with retrieved passages from your own documents so answers cite "
        "private material instead of only memorized training data.",
    )
    add_h(doc, "Why RAG exists", 2)
    add_p(
        doc,
        "Large models may hallucinate or lack up-to-date private facts. Retrieval grounds answers in snippets "
        "the user actually uploaded.",
    )
    add_h(doc, "RAG flow (boxes)", 2)
    add_code_block(
        doc,
        """User question
     |
     v
+-------------+     +------------------+     +------------------+
| Embed query | --> | Score all chunks | --> | Top-k chunks     |
| (optional)  |     | vector + lexical |     | build_context()  |
+-------------+     +------------------+     +---------+--------+
                                                      |
                                                      v
                                           +---------------------+
                                           | Prompt = context +  |
                                           | question -> Gemini  |
                                           +---------------------+
                                                      |
                                                      v
                                           Assistant reply + citations
""",
    )
    add_h(doc, "Pipeline mapped to this project", 2)
    add_p(
        doc,
        "1) Document loading: upload saves bytes; document_service.extract_text produces a single string. "
        "2) Chunking: rag_service.chunk_text splits normalized whitespace text into overlapping windows "
        "(default size 900 chars, overlap 180). "
        "3) Embedding: _embed_text posts each chunk to Gemini text-embedding-004; vectors stored in MongoDB "
        "document_chunks.embedding. "
        "4) Vector store: MongoDB collection document_chunks (not a dedicated vector DB). "
        "5) Retrieval: retrieve loads user chunks, scores 0.65 * cosine + 0.35 * token overlap. "
        "6) LLM: legal_chat sends prompt with context slice to Gemini generateContent. "
        "7) Response: ChatMessageResponse returns assistant text and citations built from chunk excerpts.",
    )
    add_h(doc, "Embeddings explained", 2)
    add_p(
        doc,
        "An embedding is a list of numbers (e.g., hundreds of floats) representing semantic meaning. "
        "Similar meanings yield vectors that point in similar directions; cosine similarity measures the angle between vectors.",
    )
    add_h(doc, "Vector database and similarity", 2)
    add_p(
        doc,
        "Dedicated vector databases optimize approximate nearest neighbor search. Here, vectors live in MongoDB "
        "documents and retrieval loops all user chunks (up to 3000), computing cosine in Python. This is simple but "
        "may be slower at large scale.",
    )
    add_h(doc, "Chunks: size trade-offs", 2)
    add_p(
        doc,
        "Too-large chunks dilute relevance and waste context window; too-small chunks lose surrounding legal meaning. "
        "Overlap helps preserve continuity across boundaries.",
    )
    add_h(doc, "Code walkthrough — RAG functions", 2)
    add_table(
        doc,
        ["Step", "File / symbol", "Role"],
        [
            ("Chunking", "rag_service.chunk_text", "Regex normalize whitespace; sliding windows with overlap"),
            ("Embed chunk", "rag_service._embed_text", "POST embedContent; returns None without API key"),
            ("Index", "rag_service.index_document_chunks", "Deletes prior chunks for doc; insert_many with embedding and token_set"),
            ("Retrieve", "rag_service.retrieve", "Hybrid vector + lexical score; sort descending"),
            ("Context", "rag_service.build_context", "Concatenate labeled chunks up to max_chars"),
            ("Citations", "rag_service.build_citations", "File name + chunk index + excerpt"),
            ("Use in chat", "chat/routes.send_message", "Calls retrieve, build_context, ai_service.legal_chat"),
        ],
    )
    add_h(doc, "RAG interview Q&A (10+)", 2)
    rag_qs = [
        (
            "What is the retrieval scoring formula in this codebase?",
            "score = 0.65 * cosine_similarity(embedding) + 0.35 * (intersection of query/chunk tokens / query token count).",
        ),
        (
            "What embedding model is used?",
            "Google Gemini text-embedding-004 via embedContent REST endpoint.",
        ),
        (
            "Is LangChain used?",
            "No. Retrieval and prompting are implemented directly in Python services.",
        ),
        (
            "What happens if embeddings are unavailable?",
            "vector term is 0; lexical token overlap can still rank chunks; if no chunks, context is empty.",
        ),
        (
            "Why store token_set in MongoDB?",
            "To speed lexical overlap using pre-tokenized words of length >= 3 characters.",
        ),
        (
            "How are citations produced?",
            "build_citations lists file_name and first ~220 characters of each retrieved chunk.",
        ),
        (
            "When are chunks (re)indexed?",
            "On upload after extract_text, and again after analyze completes to refresh text-derived chunks.",
        ),
        (
            "Does this use pgvector or Chroma?",
            "No. Embeddings are stored as arrays on MongoDB documents; similarity is computed in application code.",
        ),
        (
            "What limits retrieval fan-out?",
            "retrieve fetches up to 3000 chunk documents for the user from MongoDB.",
        ),
        (
            "How does the prompt instruct the model?",
            "legal_chat prepends Indian-law assistant instructions and includes trimmed context up to ~4000 chars.",
        ),
        (
            "What ethical caveat appears in prompts?",
            "The prompt states educational guidance, not legal advice.",
        ),
    ]
    for q, a in rag_qs:
        doc.add_paragraph(q, style="List Number")
        doc.add_paragraph(a)

    page_break(doc)

    # ----- SECTION 4 Python concepts -----
    add_h(doc, "Section 4: Python Concepts Used", 1)
    concepts = [
        (
            "Type hints (PEP 484)",
            "Annotate variables and returns like user_id: str for readability and tooling.",
            "deps.py: async def get_current_user_id(...) -> str",
            [
                ("Why use | None?", "Union types express optional values clearly in Python 3.10+."),
                ("Do hints enforce at runtime?", "Not by default; Pydantic validates API layers."),
            ],
        ),
        (
            "async / await",
            "Allows non-blocking I/O in asynchronous functions.",
            "All Mongo operations use await; httpx AsyncClient in AI calls.",
            [
                ("What runs the event loop?", "Uvicorn drives ASGI apps; await yields control while waiting on I/O."),
            ],
        ),
        (
            "Dependency injection with FastAPI Depends",
            "Reusable dependencies injected into route parameters.",
            "get_current_user_id and get_db pattern on routes.",
            [
                ("Benefit?", "Centralizes auth and DB access, fewer duplicated lines."),
            ],
        ),
        (
            "Comprehensions and generators",
            "Compact loops building lists or dicts.",
            "List comprehensions when mapping Mongo docs to Pydantic responses.",
            [
                ("Example?", "[_map_doc(d) for d in docs] in documents routes."),
            ],
        ),
        (
            "Pathlib",
            "Object-oriented filesystem paths.",
            "document_service uses Path(settings.upload_dir) / file_id.",
            [
                ("Why Path over strings?", "Cross-platform joins and readability."),
            ],
        ),
        (
            "Regular expressions (re)",
            "Pattern matching on strings for risk detection and tokenization.",
            "rag_service._tokenize and AIService.detect_risks patterns.",
            [
                ("Flag used?", "Case-insensitive search uses lowered text or inline patterns."),
            ],
        ),
        (
            "Classes and singleton service instances",
            "AIService class instantiated once as ai_service module global.",
            "Shared stateless service object reused across requests.",
            [
                ("Is this thread-safe?", "For this MVP, yes—no mutable shared fields beyond logger."),
            ],
        ),
    ]
    for title, what, where, extra_q in concepts:
        add_h(doc, title, 2)
        add_p(doc, "Concept: " + what)
        add_p(doc, "In this project: " + where)
        for q in extra_q:
            doc.add_paragraph(q[0], style="List Bullet")
            doc.add_paragraph(q[1])

    page_break(doc)

    # ----- SECTION 5 walkthrough -----
    add_h(doc, "Section 5: Complete Code Walkthrough", 1)
    add_p(
        doc,
        "This section walks repository Python modules and summarizes each function. Frontend TSX files follow.",
    )

    py_walk = [
        (
            "backend/app/main.py",
            [
                ("startup", "None", "await connect_to_mongo(); ensure_indexes", "None"),
                ("shutdown", "None", "close Mongo client", "None"),
                ("health", "None", "JSON status ok", "dict"),
            ],
        ),
        (
            "backend/app/core/config.py",
            [
                ("Settings", "env vars", "Loads configuration via pydantic-settings", "settings singleton"),
            ],
        ),
        (
            "backend/app/core/deps.py",
            [
                ("get_db", "None", "Returns AsyncIOMotorDatabase", "database"),
                ("get_current_user_id", "Bearer token", "Validates access JWT, returns sub", "str user id"),
            ],
        ),
        (
            "backend/app/core/security.py",
            [
                ("hash_password", "plain password", "pbkdf2_sha256 hash", "str"),
                ("verify_password", "plain, hash", "boolean verify", "bool"),
                ("create_token", "subject, type, minutes", "JWT encode with exp", "str jwt"),
                ("decode_token", "jwt string", "decode with secret", "payload dict"),
            ],
        ),
        (
            "backend/app/db/mongo.py",
            [
                ("connect_to_mongo", "None", "Initializes global client/db", "None"),
                ("close_mongo_connection", "None", "Closes client", "None"),
                ("get_database", "None", "Returns db or raises", "AsyncIOMotorDatabase"),
            ],
        ),
        (
            "backend/app/db/indexes.py",
            [
                ("ensure_indexes", "db", "create_index on collections", "None"),
            ],
        ),
        (
            "backend/app/modules/auth/routes.py",
            [
                ("register", "UserCreate", "Insert user if email free", "UserResponse"),
                ("login", "UserLogin", "Verify password; issue tokens", "TokenResponse"),
                ("refresh_token", "dict with refresh_token", "Validate refresh type; new tokens", "TokenResponse"),
                ("me", "Authorization", "Fetch user by id", "UserResponse"),
                ("update_me", "UserUpdate", "PATCH name", "UserResponse"),
            ],
        ),
        (
            "backend/app/modules/documents/routes.py",
            [
                ("upload_document", "UploadFile", "Save file, extract text, insert doc, index chunks, notify", "DocumentResponse"),
                ("list_documents", "auth", "List user docs desc", "list DocumentResponse"),
                ("get_document", "id", "Single doc", "DocumentResponse"),
                ("delete_document", "id", "Cascade delete related rows", "ok flag"),
                ("analyze_document", "id", "Heuristic clauses/risks/summary; update risk_score", "{\"ok\": true}"),
            ],
        ),
        (
            "backend/app/modules/analysis/routes.py",
            [
                ("list_clauses", "document_id", "Requires ownership", "ClauseResponse list"),
                ("list_risks", "document_id", "Requires ownership", "RiskResponse list"),
                ("get_summary", "document_id", "Summary record", "SummaryResponse"),
                ("explain_clause", "ClauseExplainRequest", "Calls ai_service.explain_clause", "dict"),
            ],
        ),
        (
            "backend/app/modules/chat/routes.py",
            [
                ("create_chat_session", "ChatCreate", "Insert chat shell", "id + title"),
                ("list_sessions", "auth", "Recent chats", "list metadata"),
                ("get_session", "session_id", "Messages array", "dict"),
                ("send_message", "ChatMessageRequest", "RAG retrieve + legal_chat; persist messages", "ChatMessageResponse"),
            ],
        ),
        (
            "backend/app/modules/search/routes.py",
            [
                ("search_all", "query string q", "Regex searches across filenames, titles, clauses, risks", "JSON buckets"),
            ],
        ),
        (
            "backend/app/modules/notifications/routes.py",
            [
                ("list_notifications", "auth", "Recent notifications", "list dict"),
            ],
        ),
        (
            "backend/app/services/document_service.py",
            [
                ("save_upload", "UploadFile", "Write bytes to uploads/uuid.ext", "path, ext"),
                ("extract_text", "path, ext", "PDF/DOCX text or image placeholder", "str"),
            ],
        ),
        (
            "backend/app/services/rag_service.py",
            [
                ("chunk_text", "text, sizes", "Whitespace normalize + overlap slicing", "list[str]"),
                ("_tokenize", "text", "Word tokens length>=3", "set[str]"),
                ("_embed_text", "text", "Gemini embed", "list[float]|None"),
                ("_cosine", "vectors", "Cosine similarity", "float"),
                ("index_document_chunks", "db + ids + text", "Replace chunks for doc", "int count"),
                ("retrieve", "db + user + query", "Hybrid score top_k", "list[dict]"),
                ("build_context", "chunks", "Join with labels", "str"),
                ("build_citations", "chunks", "Citation dicts", "list[dict]"),
            ],
        ),
        (
            "backend/app/services/ai_service.py",
            [
                ("_call_gemini", "prompt", "POST generateContent", "str|None"),
                ("extract_clauses", "full text", "Keyword buckets -> pseudo clauses", "list[dict]"),
                ("detect_risks", "clauses + text", "Regex library across corpus", "list[dict]"),
                ("summarize_contract", "text", "First sentences heuristics", "dict"),
                ("legal_chat", "question, context", "Gemini or fallback answer + citations structure", "dict"),
                ("explain_clause", "type, content", "Static explanatory template", "dict"),
            ],
        ),
        (
            "backend/tests/test_health.py",
            [
                ("test_health", "uses TestClient", "GET /health returns ok", "asserts"),
            ],
        ),
    ]

    for file, rows in py_walk:
        add_h(doc, file, 2)
        add_table(doc, ["Symbol", "Inputs", "Behavior", "Returns"], list(rows))

    add_h(doc, "Schemas (Pydantic)", 2)
    add_table(
        doc,
        ["File", "Models", "Purpose"],
        [
            ("schemas/common.py", "UserCreate, UserLogin, TokenResponse, UserResponse, UserUpdate, Message", "Auth and user DTOs"),
            (
                "schemas/legal.py",
                "DocumentResponse, ClauseResponse, RiskResponse, SummaryResponse, ClauseExplainRequest, ChatCreate, ChatMessageRequest, Citation, ChatMessageResponse",
                "Legal/chat API contracts with Literal clause types",
            ),
        ],
    )
    add_h(doc, "Frontend (TypeScript / React) file-by-file", 2)
    fe = [
        ("src/main.tsx", "Bootstraps React with StrictMode, QueryClientProvider, BrowserRouter, Toaster."),
        ("src/App.tsx", "AuthProvider; routes /, /login, /register, /app/* with ProtectedRoutes guarding workspace."),
        ("src/lib/api.ts", "Axios instance with VITE_API_BASE_URL; attaches Bearer; refreshAccessToken posts /auth/refresh."),
        ("src/lib/types.ts", "User, DocumentItem, risk and analysis status unions."),
        ("src/lib/utils.ts", "cn() merges Tailwind classes via clsx + tailwind-merge."),
        ("src/hooks/use-auth.tsx", "AuthContext with user, loading, setUser, logout; boot fetch /auth/me with refresh retry."),
        ("src/pages/auth-pages.tsx", "LoginPage and RegisterPage forms posting to /auth/*; stores tokens; navigates to /app/dashboard."),
        ("src/pages/workspace-pages.tsx", "Dashboard, Chat (sessions, upload-in-chat, messages), Upload, Documents (clauses/risks/summary), Analysis stub, Profile, Settings."),
        ("src/pages/landing-page.tsx", "Marketing landing with animated feature carousel, how-it-works, use cases, testimonials, FAQ, pricing CTA."),
        ("src/components/app-sidebar.tsx", "Legalyze-branded nav links to workspace routes using useLocation active styles."),
        ("src/components/chat-bubble.tsx", "User vs assistant layout; ReactMarkdown body; citations block; copy button."),
        ("src/components/upload-zone.tsx", "Drag/drop or pick file UI used by UploadPage."),
        ("src/components/feature-card.tsx", "Reusable marketing card if referenced from landing."),
        ("src/components/marketing-navbar.tsx", "Top nav for landing linking sections and auth."),
        ("src/components/risk-badge.tsx", "Visual severity badge for numeric or categorical risk."),
        ("src/components/ui/button.tsx, card.tsx, input.tsx, skeleton.tsx", "Styled primitives (likely shadcn-style) for consistent UI."),
        ("src/index.css, App.css", "Global and app-level styles; Tailwind layer."),
        ("public/, index.html", "Static favicon/icons and Vite HTML shell."),
    ]
    add_table(doc, ["Path", "Role"], fe)
    add_p(doc, "Shared folder: shared/contracts/legal.ts exports clauseTypes array and ClauseType union mirroring backend literals.")
    add_h(doc, "Design decisions", 2)
    add_p(
        doc,
        "- Mongo stores chat messages inline arrays for MVP simplicity (not a separate messages collection). "
        "- Risk scoring sums severities with caps at 100. "
        "- Auth tokens in login use literal 30 minutes and 7 days rather than reading settings.access_token_exp_minutes—"
        "an inconsistency to fix in production hardening. "
        "- CORS allow_origins=['*'] is permissive for development only. "
        "- Image uploads do not perform real OCR yet; placeholder text reminds integrators to add vision/OCR. "
        "- explain_clause is template-based, not LLM-driven, keeping latency and cost low.",
    )
    page_break(doc)

    # ----- SECTION 6 API -----
    add_h(doc, "Section 6: API / Endpoints", 1)
    add_table(
        doc,
        ["Method", "Path", "Auth", "Description"],
        [
            ("GET", "/health", "No", "Liveness check"),
            ("POST", "/auth/register", "No", "Create user"),
            ("POST", "/auth/login", "No", "Issue JWT pair"),
            ("POST", "/auth/refresh", "No", "Refresh tokens (body.refresh_token)"),
            ("GET", "/auth/me", "Bearer", "Current profile"),
            ("PATCH", "/users/me", "Bearer", "Update profile name"),
            ("POST", "/documents/upload", "Bearer multipart", "Upload + index chunks"),
            ("GET", "/documents", "Bearer", "List documents"),
            ("GET", "/documents/{id}", "Bearer", "Get metadata"),
            ("DELETE", "/documents/{id}", "Bearer", "Delete doc + dependent records"),
            ("POST", "/documents/{id}/analyze", "Bearer", "Run heuristics analysis"),
            ("GET", "/documents/{id}/clauses", "Bearer", "List clauses"),
            ("GET", "/documents/{id}/risks", "Bearer", "List risks"),
            ("GET", "/documents/{id}/summary", "Bearer", "Summary"),
            ("POST", "/documents/{id}/clauses/explain", "Bearer JSON", "Explain clause"),
            ("POST", "/chat/sessions", "Bearer", "Create chat"),
            ("GET", "/chat/sessions", "Bearer", "List chats"),
            ("GET", "/chat/sessions/{id}", "Bearer", "Fetch messages"),
            ("POST", "/chat/sessions/{id}/messages", "Bearer", "Send user message (RAG)"),
            ("GET", "/search?q=", "Bearer", "Regex search"),
            ("GET", "/notifications", "Bearer", "Notifications"),
        ],
    )
    add_h(doc, "Example requests / responses", 2)
    add_code_block(
        doc,
        """POST /auth/login
Request JSON: {\"email\":\"user@example.com\",\"password\":\"********\"}
Response: {\"access_token\":\"...\",\"refresh_token\":\"...\",\"token_type\":\"bearer\"}

POST /documents/upload
Headers: Authorization: Bearer <access>
Body: multipart/form-data file=@contract.pdf
Response: {\"id\":\"...\",\"file_name\":\"contract.pdf\",\"file_type\":\"pdf\",\"analysis_status\":\"uploaded\",...}

POST /chat/sessions/{id}/messages
Request: {\"content\":\"What are termination rights?\"}
Response: {\"session_id\":\"...\",\"user_message\":\"...\",\"assistant_message\":\"...\",\"citations\":[{\"source\":\"...\",\"excerpt\":\"...\"}]}
""",
    )
    add_h(doc, "REST / FastAPI interview Q&A", 2)
    for q, a in [
        (
            "Why separate routers with prefixes?",
            "Namespaces endpoints (/auth vs /documents) and attaches shared tags for docs.",
        ),
        (
            "How is authentication enforced?",
            "Depends(get_current_user_id) ensures Bearer JWT with type access.",
        ),
        (
            "What status codes are used for auth failures?",
            "401 for invalid token/credentials, 404 when documents missing.",
        ),
    ]:
        doc.add_paragraph(q, style="List Number")
        doc.add_paragraph(a)

    page_break(doc)

    # ----- SECTION 7 data flow -----
    add_h(doc, "Section 7: Data Flow Diagram", 1)
    flows = [
        "User opens React app; AuthProvider calls GET /auth/me with stored access token.",
        "Login posts credentials; tokens saved to localStorage; axios interceptor attaches Authorization header.",
        "Upload: browser sends multipart to POST /documents/upload; backend saves file, extracts text, inserts Mongo document row, calls rag_service.index_document_chunks, inserts notification.",
        "Analyze: POST /documents/{id}/analyze loads extracted text, runs ai_service heuristics, writes clauses/risks/summaries, updates analysis_status and risk_score, reindexes chunks, notifies user.",
        "Chat message: POST /chat/sessions/{id}/messages loads session, rag_service.retrieve scores chunks, build_context, ai_service.legal_chat calls Gemini or fallback, persists user+assistant messages.",
        "Search: GET /search applies Mongo $regex queries across permitted collections scoped to user.",
    ]
    for i, step in enumerate(flows, 1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")

    page_break(doc)

    # ----- SECTION 8 setup -----
    add_h(doc, "Section 8: Setup & Configuration", 1)
    add_h(doc, "Install and run from zero", 2)
    add_p(doc, "Docker: docker compose up --build — launches mongo:7, backend image (uvicorn), frontend image (nginx static).")
    add_p(
        doc,
        "Local backend: python -m venv .venv; activate; pip install -r requirements.txt; "
        "copy .env.example to .env; run uvicorn app.main:app --reload inside backend/.",
    )
    add_p(
        doc,
        "Local frontend: npm install; npm run dev; set frontend/.env VITE_API_BASE_URL to backend URL.",
    )
    add_h(doc, "Environment variables", 2)
    add_table(
        doc,
        ["Variable", "Purpose"],
        [
            ("APP_NAME", "FastAPI title / branding"),
            ("ENVIRONMENT", "Deployment stage string"),
            ("MONGO_URI", "Mongo connection URI (docker-compose uses mongodb://mongo:27017)"),
            ("MONGO_DB_NAME", "Database name (default legal_ai)"),
            ("JWT_SECRET", "Signing key for HS256 tokens—must be strong in production"),
            ("JWT_ALGORITHM", "Usually HS256"),
            ("ACCESS_TOKEN_EXP_MINUTES", "Intended access token lifetime (note: code may still use literals in auth routes)"),
            ("REFRESH_TOKEN_EXP_MINUTES", "Refresh lifetime"),
            ("UPLOAD_DIR", "Directory for stored uploads"),
            ("GEMINI_API_KEY", "Google AI Studio key for LLM + embeddings"),
            ("GEMINI_MODEL", "Model id for generateContent (default gemini-1.5-flash)"),
            ("VITE_API_BASE_URL", "Frontend axios base URL"),
        ],
    )
    add_h(doc, "Common errors", 2)
    for q, a in [
        ("Mongo connection refused", "Ensure Mongo is running and MONGO_URI matches host/port; docker service name is mongo."),
        ("401 on protected routes", "Token expired—use refresh flow or login again; check Authorization header."),
        ("Empty PDF text", "Scanned PDFs need OCR; pipeline returns little text so analysis may be sparse."),
        ("Gemini errors", "Verify GEMINI_API_KEY; watch rate limits—service retries once on some HTTP errors."),
    ]:
        doc.add_paragraph(q, style="List Bullet")
        doc.add_paragraph(a)

    page_break(doc)

    # ----- SECTION 9 master interview -----
    add_h(doc, "Section 9: Master Interview Preparation", 1)
    master = [
        # Python
        ("Python Basics", "What is the difference between a list and a tuple?", "Lists are mutable; tuples immutable—use tuples for fixed records."),
        ("Python Basics", "What does async def accomplish in FastAPI?", "Declares a coroutine route that can await I/O without blocking other requests."),
        ("Python Basics", "How are dictionaries used when returning Mongo documents?", "Keys like _id map to fields; ObjectId may be converted to str for JSON."),
        ("Python Basics", "What is Pydantic?", "Data validation layer—defines schemas for request/response models."),
        ("Python Basics", "Explain try/except around password hashing in register.", "Turns low-level hashing errors into HTTP 400 with safe messaging."),
        # RAG
        ("RAG & LLMs", "Define retrieval-augmented generation.", "Fetch relevant texts first, then condition the LLM prompt on them."),
        ("RAG & LLMs", "Why cosine similarity?", "Measures directional alignment of embedding vectors independent of magnitude scaling when norms considered."),
        ("RAG & LLMs", "What is a chunk overlap?", "Characters repeated between consecutive chunks to preserve cross-boundary context."),
        ("RAG & LLMs", "How does this app ground answers?", "Inserts top chunks into the Gemini prompt via legal_chat."),
        ("RAG & LLMs", "Fallback when LLM absent?", "Rule-based messages in _fallback_legal_chat_answer."),
        # Vector
        ("Vector Databases", "Why might MongoDB + Python cosine be a bottleneck?", "Linear scan over user chunks; specialized ANN indexes scale better."),
        ("Vector Databases", "What is ANN?", "Approximate nearest neighbor search trading accuracy for speed."),
        ("Vector Databases", "Does this project normalize embeddings before cosine?", "Cosine function divides by norms of both vectors (implicit normalization effect)."),
        # API
        ("API Development", "How does FastAPI validate an email field?", "EmailStr from Pydantic in UserCreate."),
        ("API Development", "Why use response_model?", "Filters output fields and documents schema for clients."),
        ("API Development", "How to secure cookies vs tokens here?", "Bearer tokens in headers—cookies not used in this MVP."),
        # System design
        ("System Design", "How would you shard chat history at scale?", "Move messages to their own collection with indexed session_id."),
        ("System Design", "How to improve OCR?", "Integrate Gemini vision or Tesseract; preprocess PDF images."),
        ("System Design", "How to harden auth?", "Align token TTL with settings, rotate secrets, narrow CORS, rate limit login."),
        ("System Design", "How to scale retrieval?", "Use a vector database with ANN; pre-filter by user_id metadata."),
        ("Improvements", "Add integration tests for auth and RAG.", "Use TestClient with seeded Mongo or mongomock."),
        ("Improvements", "Add structured logging/trace ids.", "Observability across Gemini and Mongo calls."),
        ("Improvements", "Add file virus scanning.", "ClamAV or cloud malware API before saving uploads."),
        ("Improvements", "Implement background jobs for analyze.", "Celery/RQ so HTTP returns quickly."),
        ("Improvements", "Internationalization on frontend.", "i18n framework with extracted strings."),
        ("Improvements", "Role-based access control.", "Tenant or role claims inside JWT for team workspaces."),
        ("Python Basics", "What is a list comprehension?", "Compact syntax [f(x) for x in items if cond]—used mapping Mongo docs."),
        ("Python Basics", "What does Pathlib provide?", "OO path operations—used when saving uploads."),
        ("RAG & LLMs", "Why include lexical scoring?", "Captures keyword overlap when embeddings missing or weak."),
        ("RAG & LLMs", "Prompt injection concern?", "Untrusted documents included in prompts—sanitize or isolate untrusted content."),
        ("Vector Databases", "Difference between exact and approximate search?", "Exact checks all pairs; approximate uses indexes/graphs for speed."),
        ("API Development", "What is multipart upload?", "HTTP encoding for files—FastAPI UploadFile handles streams."),
        ("API Development", "Why HTTPBearer?", "Standard Authorization: Bearer header parsing."),
        ("System Design", "Caching strategies for dashboard counts?", "Redis or memoized queries with ttl."),
        ("System Design", "How to version API?", "Prefix routes /v1 and maintain backward compatible schemas."),
        ("Improvements", "Add LLM-based clause extraction.", "Replace keyword method with structured JSON outputs validated by Pydantic."),
        ("Improvements", "Per-tenant encryption.", "KMS-managed keys for documents at rest."),
        ("Python Basics", "What is a module-level singleton?", "ai_service = AIService() reused across imports."),
        ("Python Basics", "Why datetime timezone.utc?", "Store timestamps timezone-aware to avoid DST ambiguity."),
        ("RAG & LLMs", "Temperature not shown—why?", "Gemini REST payload here only sets text parts; tuning omitted."),
        ("RAG & LLMs", "Context truncation at 4000 chars—why?", "Balances prompt size vs latency/cost."),
        ("Vector Databases", "Embedding dimensionality?", "Defined by model—values list length must match for cosine."),
        ("API Development", "How to document multipart endpoints?", "FastAPI infers file parameters; appears in OpenAPI as multipart."),
        ("System Design", "Zero-downtime deploy?", "Blue/green behind load balancer; migrate Mongo with compatible schema."),
        ("System Design", "Backpressure on Gemini?", "Queue analyze jobs; limit concurrent httpx calls."),
        ("Improvements", "Add pytest coverage for RAG scoring.", "Golden vectors with known cosine outcomes."),
        ("Improvements", "Feature flags for Gemini.", "Toggle LLM per tenant in settings."),
    ]
    extra_master = [
        ("Python Basics", "What is ObjectId in Mongo?", "BSON type for primary keys; converted to str for API responses."),
        ("Python Basics", "Why use Field(min_length=8) on passwords?", "Enforces minimum complexity at validation layer."),
        ("RAG & LLMs", "What is a fallback path?", "Code path when external LLM is unavailable to keep UX working."),
        ("RAG & LLMs", "Why label chunks with file_name?", "Citations remain human-readable in the UI."),
        ("Vector Databases", "What is metadata filtering?", "Restricting vector search by user_id before distance calc—done here by query filter."),
        ("Vector Databases", "Why store embeddings in-app vs recompute?", "Trade storage for latency; recomputing each time would be slower and costlier."),
        ("API Development", "What is idempotency?", "Repeating a request has same effect; DELETE is idempotent; analyze POST is not strictly idempotent."),
        ("API Development", "Why return 404 vs 403?", "This MVP hides existence of other users' docs by 404 on mismatch."),
        ("System Design", "How to audit AI outputs?", "Log prompts, model version, and outputs with user id for compliance."),
        ("System Design", "Data residency considerations?", "Host Mongo and call regional Gemini endpoints as required."),
        ("Improvements", "Add contract templates library.", "Pre-seed chunk store with standard clauses for comparison."),
        ("Python Basics", "What is * and ** unpacking?", "Not heavily used here, but dict payloads use explicit keys for Mongo updates."),
        ("RAG & LLMs", "Why cap context to 5000 chars in build_context?", "Prevents overlong prompts and controls cost."),
        ("API Development", "What is OpenAPI tag?", "Groups endpoints in docs—see tags in APIRouter declarations."),
    ]
    master.extend(extra_master)
    topics: dict[str, list[tuple[str, str]]] = {}
    for topic, q, a in master:
        topics.setdefault(topic, []).append((q, a))

    for topic, qa_list in topics.items():
        add_h(doc, topic, 2)
        for q, a in qa_list:
            doc.add_paragraph(q, style="List Number")
            doc.add_paragraph(a)

    add_p(doc, f"Total curated Q&A items in this section: {len(master)}.")

    page_break(doc)

    add_h(doc, "Section 10: MongoDB Collections & Fields", 1)
    add_p(
        doc,
        "MongoDB groups records into collections (like flexible tables). The app infers schemas from Python "
        "dictionaries inserted at runtime; indexed fields accelerate common queries (see backend/app/db/indexes.py).",
    )
    add_table(
        doc,
        ["Collection", "Purpose", "Representative fields (from routes/services)"],
        [
            ("users", "Registered accounts.", "email (unique), password_hash (passlib), name, _id."),
            ("documents", "One row per uploaded file.", "user_id, file_name, file_type, storage_path, extracted_text, analysis_status, risk_score (0–100 cap), created_at."),
            ("clauses", "Heuristic clause snippets after analyze.", "document_id, clause_type, content, explanation, risk_level, created_at."),
            ("risks", "Regex/heuristic flags tied to clauses.", "document_id, clause_id?, clause_type, severity (high/medium/low), issue, risky_text, why_risky, recommendation, created_at."),
            ("summaries", "summary dict from summarize_contract.", "document_id, plain_summary, obligations[], deadlines[], payment_terms[], termination_conditions[], key_risks[], created_at."),
            ("document_chunks", "RAG store per user/document.", "user_id, document_id, file_name, chunk_index, content, token_set[], embedding[] (Gemini dims) | null, created_at."),
            ("chats", "Chat threads; messages nested array.", "_id as session id, user_id, title?, messages[{role, content, citations?, created_at}], updated_at."),
            ("notifications", "Toast-style events.", "user_id, title, message, created_at."),
        ],
    )
    add_h(doc, "Cascade & consistency notes", 2)
    add_p(
        doc,
        "Deleting a document (DELETE /documents/{id}) removes related clauses, risks, summaries, and document_chunks "
        "for that document_id. Analyze replaces prior clauses/risks/summary for the document."
    )
    add_h(doc, "Risk score calculation (interview)", 2)
    add_p(
        doc,
        "In analyze_document, each risk increments risk_score by 80 if severity is high, 50 if medium, 20 if low; "
        "final score stored as min(sum, 100). This is a simple additive heuristic."
    )

    page_break(doc)

    add_h(doc, "Section 11: Beginner Concepts (read this first if new to Python or web backends)", 1)
    add_h(doc, "How software in this repo fits together", 2)
    add_p(
        doc,
        "Your browser loads JavaScript bundles built from the React frontend. When you submit a form or click Chat, "
        "the frontend sends an HTTP request to the backend URL (often http://localhost:8000 during development). "
        "FastAPI inspects headers and JSON or multipart bodies, validates them with Pydantic models, then runs Python "
        "functions that query MongoDB (via Motor) or call Gemini (via httpx). The response travels back as JSON and "
        "React redraws components."
    )
    add_table(
        doc,
        ["Idea", "Plain explanation", "Where you see it here"],
        [
            ("Python module", "A .py file you import elsewhere.", "app.services.rag_service imported as rag_service."),
            ("Virtual environment (.venv)", "Isolated dependency folder so pip packages do not clash.", "Recommended in README before pip install."),
            ("async def / await", "Pauses only the current request while waiting on network/database.", "Nearly every Mongo call uses await."),
            ("JSON", "Structured text exchanged over HTTP ({\"key\":\"value\"}).", "Auth and chat payloads."),
            ("Environment variables", "Config outside Git; Production secrets.", ".env copied from .env.example."),
            ("REST resource", "A URL noun like /documents with verbs GET/POST/DELETE.", "routers prefixed /documents."),
            ("JWT split", "Access token proves identity for minutes; refresh token renews quietly.", "/auth/login + /auth/refresh."),
        ],
    )
    add_h(doc, "What “embedding” and “dimension” mean (no math prerequisites)", 2)
    add_p(
        doc,
        "Think of embedding as turning a paragraph into hundreds of sliders (numbers); similar paragraphs yield similar slider positions. "
        "Cosine similarity compares two sliders’ angles: 1 means very aligned, 0 means unrelated."
    )

    page_break(doc)

    add_h(doc, "Section 12: Comparing Ways to Teach the Model About Your Documents", 1)
    add_table(
        doc,
        ["Approach", "How it works", "Pros", "Cons", "This repo"],
        [
            ("Naive long prompt", "Paste entire contract into Gemini each time.", "Simple mental model.", "Hits token limits quickly; pricey; brittle.", "Legal chat trims context (~4000 chars) + top chunks—not whole doc blindly."),
            ("Fine-tuning", "Train new weights with domain data.", "Learns stylistic quirks.", "Expensive pipelines; staleness unless retrained.", "Not used."),
            ("Classic RAG (this codebase)", "Chunk, embed/stash, retrieve, answer with citations-ish snippets.", "Updates when user uploads docs; grounding.", "Retriever quality bottleneck; hallucination risk if chunks wrong.", "Yes—Gemini embeddings + lexical hybrid."),
            ("External vector DB (Pinecone, etc.)", "Same as RAG but ANN indexing.", "Faster retrieval at billions of vectors.", "Extra infra complexity.", "Not yet—vectors live in Mongo + Python cosine."),
        ],
    )
    add_h(doc, "Interview framing", 2)
    add_p(
        doc,
        "Be ready to say: embeddings give semantic fuzziness, lexical fallback keeps keyword-heavy legal terms relevant when "
        "API keys fail, Mongo gives flexible schema plus single DB story, drawbacks are brute-force scoring (3000 chunks). "
        "Production would add ANN and background workers."
    )

    page_break(doc)

    add_h(doc, "Section 13: Security, Privacy & Limitations", 1)
    add_table(
        doc,
        ["Topic", "Current behavior", "Hardening idea"],
        [
            ("CORS", "allow_origins=['*'] in backend/app/main.py (dev convenience).", "Restrict to frontend origin."),
            ("Tokens", "Bearer stored in browser localStorage (common MVP).", "Consider httpOnly cookie + CSRF defenses."),
            ("Prompt injection", "Uploaded doc text appended to Gemini prompt.", "Strip instructions, classify untrusted blobs."),
            ("OCR gaps", "Image uploads return placeholder string only.", "Add vision API or OCR service."),
            ("Legal disclaimer", "Prompt says guidance not advice.", "Add UI banner + jurisdiction-specific copy."),
            ("Secrets", "JWT secret env driven but default weak in dev.", "Rotate keys; KMS in prod."),
            ("Rate limiting", "Not implemented on auth routes.", "Add slowapi or API gateway quotas."),
            ("File malware", "No AV scan.", "Hook ClamAV or cloud scanner."),
        ],
    )

    page_break(doc)

    add_h(doc, "Section 14: Glossary", 1)
    glossary = [
        ("API", "Application Programming Interface: how programs talk over HTTP with structured requests and responses."),
        ("ASGI", "Asynchronous Server Gateway Interface for Python web servers and apps."),
        ("Bearer token", "Authorization scheme where the secret token is placed after the word Bearer in the header."),
        ("BSON", "Binary JSON-like encoding MongoDB uses internally for documents."),
        ("Chunk", "A slice of document text used for retrieval and embedding."),
        ("Collection", "MongoDB grouping of documents, similar to a table but schema-flexible."),
        ("Context window", "Maximum prompt size a language model can consider at once."),
        ("Cosine similarity", "Similarity measure between two vectors based on the cosine of the angle between them."),
        ("CORS", "Cross-Origin Resource Sharing: browser security rules for calling APIs from different domains."),
        ("Dependency injection", "Supplying shared objects like database connections automatically to route handlers."),
        ("Docker Compose", "Tool to run multi-container applications from a YAML file."),
        ("Embedding", "Numeric vector representing text semantics for similarity comparisons."),
        ("Environment variable", "Configuration value provided outside source code, often via .env."),
        ("FastAPI", "Python framework for building APIs quickly with type hints and validation."),
        ("Gemini", "Google’s generative model family used here via REST for chat and embeddings."),
        ("Heuristic", "Rule-based approximation rather than full semantic reasoning."),
        ("HS256", "HMAC SHA-256 algorithm for signing JWTs with a shared secret."),
        ("HTTPX", "Python HTTP client supporting sync and async requests."),
        ("Index", "Database structure speeding queries—unique on email, compound on user_id + created_at, etc."),
        ("JWT", "JSON Web Token: compact signed claims for authentication."),
        ("Lexical scoring", "Scoring using overlapping words between query and passage."),
        ("LLM", "Large Language Model trained to predict and generate human-like text."),
        ("Motor", "Async Python driver for MongoDB."),
        ("MongoDB", "Document-oriented database storing BSON documents."),
        ("Multimodal", "Models handling more than text—Gemini family supports this broadly."),
        ("MVP", "Minimum viable product focusing on core features."),
        ("ObjectId", "MongoDB’s default unique document identifier type."),
        ("OCR", "Optical Character Recognition: turning images of text into text data."),
        ("OpenAPI", "Specification describing REST APIs; FastAPI generates it automatically."),
        ("Passlib", "Password hashing library—project uses pbkdf2_sha256."),
        ("PDF", "Portable Document Format; text extraction depends on whether text is embedded or scanned."),
        ("Pydantic", "Data validation using Python type annotations."),
        ("PyJWT", "Library to encode and decode JWT tokens."),
        ("RAG", "Retrieval-Augmented Generation combining search with language generation."),
        ("React", "JavaScript library for building component-based user interfaces."),
        ("Regex", "Regular expression pattern matching on strings."),
        ("Refresh token", "Longer-lived token used to obtain new access tokens."),
        ("REST", "Representational State Transfer style of HTTP APIs using resources and verbs."),
        ("Similarity search", "Finding nearest embedding vectors to a query vector."),
        ("SPA", "Single-page application where navigation occurs client-side without full page reloads."),
        ("Tailwind CSS", "Utility-first styling framework using composable class names."),
        ("Token (JWT claim)", "Typed claim distinguishing access vs refresh in custom payload field."),
        ("Token (text)", "A word or subword unit used in lexical overlap scoring."),
        ("Type hints", "Optional annotations in Python showing intended data types."),
        ("TypeScript", "JavaScript with static typing for safer front-end code."),
        ("UploadFile", "FastAPI class representing an uploaded multipart file stream."),
        ("Uvicorn", "ASGI server that runs FastAPI applications."),
        ("Vector", "List of numbers; embedding vectors live in Mongo documents."),
        ("Vite", "Fast front-end build tool and dev server used by this React project."),
    ]
    for term, defin in glossary:
        doc.add_paragraph(f"{term}: {defin}", style="List Bullet")

    add_p(doc, f"Glossary entries: {len(glossary)}.")

    out = "Legal_AI_Complete_Technical_Documentation.docx"
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
