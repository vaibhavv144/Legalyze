"""
Generate Indian Legal AI Platform complete technical documentation as .docx
Run from repo root: python scripts/generate_technical_documentation.py
Requires: pip install python-docx (listed in backend/requirements.txt)
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)


def add_mermaid_note(doc: Document, caption: str, ascii_art: str) -> None:
    doc.add_paragraph(caption, style="Heading 4")
    add_code_block(doc, ascii_art)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for r_idx, row in enumerate(rows):
        for c_idx, cell in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = cell


def main() -> None:
    root = Path(__file__).resolve().parent.parent
    out_path = root / "Indian_Legal_AI_Platform_Complete_Technical_Documentation.docx"

    doc = Document()
    doc.core_properties.title = "Indian Legal AI Platform — Complete Technical Documentation"
    doc.core_properties.subject = "Architecture, RAG, APIs, interview prep"
    doc.core_properties.keywords = "FastAPI, MongoDB, Gemini, RAG, React"

    # Title
    t = doc.add_heading("Indian Legal AI Platform", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Complete Technical Documentation").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Generated for beginners: Python, AI/RAG, APIs, databases, and this codebase."
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # --- Section 1 ---
    doc.add_heading("Section 1: Project Overview", 1)
    doc.add_paragraph(
        "This repository is a full-stack MVP for an Indian law–focused legal assistant. "
        "The backend is FastAPI with MongoDB; the frontend is React 19 with TypeScript, Vite, and Tailwind. "
        "Google Gemini (when GEMINI_API_KEY is set) powers legal chat and embeddings; without a key, "
        "the app uses rule-based fallbacks for chat and lexical-only retrieval for RAG."
    )
    doc.add_paragraph(
        "Problem solved: users need to upload contracts (PDF/DOCX/images), get clause extraction "
        "and risk heuristics, summaries, and a chat that can cite uploaded material—without running everything "
        "through a generic LLM with no document grounding."
    )
    doc.add_paragraph(
        "Target users: legal teams, startups, and individuals reviewing Indian-law-related contracts "
        "who want a structured workspace (documents, analysis, chat, search, notifications)."
    )
    doc.add_heading("Core features", 2)
    add_table(
        doc,
        ["Feature", "Description"],
        [
            ["JWT auth", "Register, login, refresh, profile (/auth, /users)"],
            ["Document upload", "PDF, DOCX, JPG/PNG; text extraction; stored in MongoDB + disk"],
            ["RAG chat", "Chunking, optional Gemini embeddings, hybrid retrieval, context to Gemini"],
            ["Contract analysis", "Keyword clause extraction, regex risk rules, heuristic summary"],
            ["Search", "MongoDB regex across documents, chats, clauses, risks"],
            ["Notifications", "Upload and analysis completion messages"],
        ],
    )
    doc.add_paragraph()
    doc.add_heading("Why this architecture", 2)
    doc.add_paragraph(
        "FastAPI gives async I/O and automatic OpenAPI docs. MongoDB fits flexible legal documents and "
        "embedded chunk vectors without a separate schema migration per field. React + Vite keeps the UI "
        "fast to develop. Gemini is chosen as a single vendor for both generation and embeddings (text-embedding-004) "
        "to reduce integration surface; the code does not use LangChain—calls are direct HTTP with httpx."
    )
    add_mermaid_note(
        doc,
        "System overview (logical)",
        """
+------------------+     HTTPS/JSON      +------------------+
|  React (Vite)    | <-----------------> |  FastAPI         |
|  localhost:5173  |   Bearer JWT        |  :8000           |
+------------------+                     +--------+---------+
                                                |
                        +-----------------------+---------------------+
                        | MongoDB                | uploads/ (files)    |
                        | users, documents,      | PDF/DOCX/images     |
                        | chats, clauses, risks, |                     |
                        | summaries, chunks      |                     |
                        +-----------------------+---------------------+
                                                |
                        +-----------------------+
                        | Google Gemini API     |
                        | (optional)            |
                        +-----------------------+
""",
    )

    # --- Section 2 ---
    doc.add_heading("Section 2: High-Level Architecture", 1)
    doc.add_paragraph(
        "The system has three main runtime parts: browser UI, API server, and MongoDB. Optional Gemini "
        "calls add a fourth external dependency. There is no separate dedicated vector database (Pinecone, "
        "Weaviate, etc.); vectors are stored as arrays on documents in the document_chunks collection."
    )
    add_mermaid_note(
        doc,
        "Request flow (authenticated API)",
        """
Client                    FastAPI                    MongoDB / Files
  |                          |                              |
  |-- POST /documents/upload -> save file, insert doc -----|
  |                          |-- index chunks (RAG) ------->|
  |                          |<- store chunks + emb --------|
  |<- 201 DocumentResponse ---|                              |
""",
    )
    doc.add_heading("Components", 2)
    items = [
        (
            "Frontend (React + TS + Tailwind)",
            "What: SPA that calls REST APIs with axios; tokens in localStorage.",
            "Why: rich UX for dashboard, upload, chat.",
            "Alternatives: Next.js, Vue. Chosen: Vite React for speed and simplicity.",
        ),
        (
            "Backend (FastAPI)",
            "What: Python ASGI app; routes in app/modules/*/routes.py.",
            "Why: async routes, Pydantic validation, OpenAPI.",
            "Alternatives: Flask, Django. Chosen: FastAPI for async + typing.",
        ),
        (
            "Database (MongoDB via Motor)",
            "What: Document store; collections include users, documents, document_chunks, chats, etc.",
            "Why: flexible JSON documents for varying legal fields.",
            "Alternatives: PostgreSQL + JSONB. Chosen: Mongo for rapid MVP.",
        ),
        (
            '“Vector storage” (MongoDB document_chunks)',
            "What: Each chunk has content, token_set, optional embedding list.",
            "Why: avoids extra infra for MVP; hybrid score still works without embeddings.",
            "Alternatives: Pinecone, Qdrant, pgvector. Chosen: simplicity.",
        ),
        (
            "LLM (Google Gemini)",
            "What: generateContent for chat and analysis prompts; disabled gracefully without API key.",
            "Why: strong multimodal path for future OCR/vision.",
            "Alternatives: OpenAI, local Llama. Chosen: single Google stack for embed + chat.",
        ),
        (
            "Embedding model (text-embedding-004)",
            "What: REST embedContent in rag_service._embed_text.",
            "Why: semantic similarity for RAG when key present.",
            "Alternatives: OpenAI text-embedding-3, local models.",
        ),
        (
            "Retrieval (RAGService.retrieve)",
            "What: Loads user chunks, scores 0.65*cosine + 0.35*lexical overlap.",
            "Why: robust when embeddings missing (lexical still contributes).",
            "Alternatives: pure vector ANN. Chosen: hybrid for resilience.",
        ),
        (
            "Document parser (DocumentService)",
            "What: pypdf for PDF, python-docx for DOCX, placeholder for images.",
            "Why: common contract formats. OCR is explicitly a placeholder.",
            "Alternatives: Apache Tika, unstructured.io. Chosen: lightweight Python libs.",
        ),
        (
            "Chunking (RAGService.chunk_text)",
            "What: 900 char windows, 180 overlap, whitespace normalized.",
            "Why: fits context and reduces boundary loss.",
            "Alternatives: token-based splitters (tiktoken). Chosen: character windows for simplicity.",
        ),
        (
            "Prompt builder (AIService.legal_chat)",
            "What: System-style instruction + context slice (up to 4000 chars) + user question.",
            "Why: grounds answers in uploaded text when available.",
            "Alternatives: LangChain chat templates. Chosen: single f-string for clarity.",
        ),
    ]
    for title, a, b, c in items:
        doc.add_paragraph(title, style="Heading 3")
        doc.add_paragraph(f"1) What it is: {a}")
        doc.add_paragraph(f"2) Why used: {b}")
        doc.add_paragraph(f"3) Alternatives / choice: {c}")

    # --- Section 3 ---
    doc.add_heading("Section 3: Folder Structure Deep Dive", 1)
    doc.add_paragraph(
        "Repository root contains README.md, docker-compose.yml, backend/, frontend/, and (after generation) "
        "this documentation. node_modules and .venv are local dependencies and should not be edited by hand."
    )
    tree = """
Legal_AI/
  README.md                 # Human runbook: Docker + local dev
  docker-compose.yml        # mongo, backend, frontend services
  backend/
    Dockerfile
    requirements.txt        # fastapi, motor, PyJWT, pypdf, python-docx, httpx, ...
    .env.example            # MONGO_URI, JWT_*, GEMINI_*
    app/
      main.py               # FastAPI app, CORS, routers, lifespan
      core/
        config.py           # pydantic-settings from env
        deps.py             # get_db, get_current_user_id (Bearer JWT)
        security.py         # hash/verify password, create/decode JWT
      db/
        mongo.py            # Motor client singleton
        indexes.py          # create_index on collections
      schemas/
        common.py           # User*, TokenResponse
        legal.py            # Document, Clause, Risk, Chat DTOs
      modules/
        auth/routes.py      # /auth/*, /users/me
        documents/routes.py # /documents/*
        analysis/routes.py  # /documents/{id}/clauses, risks, summary, explain
        chat/routes.py      # /chat/sessions*
        search/routes.py    # /search
        notifications/routes.py
      services/
        ai_service.py       # Gemini + heuristics
        rag_service.py      # chunk, embed, retrieve, context
        document_service.py # save file, extract text
  frontend/
    package.json
    vite.config.ts (if present)
    src/
      main.tsx, App.tsx
      lib/api.ts            # axios base URL + Bearer
      hooks/use-auth.tsx
      pages/workspace-pages.tsx  # dashboard, chat, upload, documents, analysis
      components/...
"""
    add_code_block(doc, tree)

    file_notes = [
        (
            "backend/app/main.py",
            "Creates FastAPI app, CORS allow all, startup connects MongoDB and ensure_indexes, "
            "shutdown closes client. Registers routers. Removing it removes the entire API.",
        ),
        (
            "backend/app/core/config.py",
            "Settings: mongo_uri, jwt_secret, gemini_api_key, gemini_model. Loaded from .env. "
            "If removed, nothing starts with correct env.",
        ),
        (
            "backend/app/core/deps.py",
            "HTTPBearer dependency; decode_token must have type 'access'. Core auth for protected routes.",
        ),
        (
            "backend/app/core/security.py",
            "passlib pbkdf2_sha256 for passwords; PyJWT for HS256 tokens.",
        ),
        (
            "backend/app/db/mongo.py",
            "Global Motor client; get_database() must run after startup.",
        ),
        (
            "backend/app/db/indexes.py",
            "Indexes on email, document user_id+time, chunks by user+document, etc.",
        ),
        (
            "backend/app/modules/auth/routes.py",
            "register, login, refresh, /auth/me, PATCH /users/me.",
        ),
        (
            "backend/app/modules/documents/routes.py",
            "upload (extract text, insert doc, rag index, notification), list, get, delete, analyze.",
        ),
        (
            "backend/app/modules/analysis/routes.py",
            "Read clauses, risks, summary; POST explain_clause (heuristic response in ai_service).",
        ),
        (
            "backend/app/modules/chat/routes.py",
            "Chat sessions CRUD; send_message runs RAG retrieve + legal_chat + persist messages.",
        ),
        (
            "backend/app/modules/search/routes.py",
            "Case-insensitive regex search across user-scoped collections.",
        ),
        (
            "backend/app/modules/notifications/routes.py",
            "List notifications for user.",
        ),
        (
            "backend/app/services/ai_service.py",
            "Gemini generateContent; clause extraction by keywords; risk detection by regex rules; "
            "summarize_contract by sentence heuristics; legal_chat with fallback strings.",
        ),
        (
            "backend/app/services/rag_service.py",
            "chunk_text, _embed_text (Gemini), retrieve hybrid score, build_context, build_citations.",
        ),
        (
            "backend/app/services/document_service.py",
            "save_upload to uploads/, extract_text PDF/DOCX/image placeholder.",
        ),
        (
            "frontend/src/lib/api.ts",
            "Axios instance; Authorization Bearer from localStorage; refresh endpoint.",
        ),
        (
            "frontend/src/hooks/use-auth.tsx",
            "On load: GET /auth/me; on 401 try refresh; logout clears tokens.",
        ),
    ]
    for path, desc in file_notes:
        doc.add_paragraph(path, style="Heading 3")
        doc.add_paragraph(desc)

    # --- Section 4 Technology ---
    doc.add_heading("Section 4: Technology Deep Dive", 1)
    tech = [
        (
            "Python 3.12",
            "High-level language with async/await. Used for the entire backend. "
            "Advantages: rich AI/HTTP ecosystem. Disadvantages: GIL limits CPU-bound threads (use async I/O).",
        ),
        (
            "FastAPI",
            "ASGI web framework. Uses Pydantic for request/response models, auto OpenAPI. "
            "This project: async route handlers, Depends() for DB and user.",
        ),
        (
            "Motor",
            "Async MongoDB driver. Non-blocking database access in the same event loop as FastAPI.",
        ),
        (
            "MongoDB",
            "NoSQL document database. Collections map to domain objects; no SQL joins—embed or reference by id string.",
        ),
        (
            "PyJWT + passlib",
            "JWT for stateless auth; passlib for password hashing (pbkdf2_sha256 in security.py).",
        ),
        (
            "httpx",
            "Async HTTP client to call Google Generative Language REST APIs.",
        ),
        (
            "pypdf / python-docx / Pillow",
            "PDF and Word text extraction; Pillow implied for image path (OCR not implemented).",
        ),
        (
            "React 19 + TypeScript + Vite",
            "Component UI, type safety, fast dev server. react-router-dom for /app/* workspace.",
        ),
        (
            "axios + TanStack Query (listed in package.json)",
            "HTTP; Query available for data fetching patterns (workspace uses useEffect + api in places).",
        ),
        (
            "Docker & docker-compose",
            "mongo:7, backend build from backend/Dockerfile, frontend multi-stage nginx serving static dist.",
        ),
    ]
    for name, body in tech:
        doc.add_paragraph(name, style="Heading 3")
        doc.add_paragraph(body)

    doc.add_paragraph(
        "Note: LangChain does not appear in requirements.txt. RAG and LLM calls are implemented manually—"
        "this is easier to read for learning but more code to maintain than a chain abstraction."
    )

    # --- Section 5 RAG ---
    doc.add_heading("Section 5: RAG Complete Deep Dive", 1)
    doc.add_heading("What is RAG?", 2)
    doc.add_paragraph(
        "Retrieval-Augmented Generation: instead of the model memorizing your contract, you store the contract "
        "in chunks, find the chunks most relevant to the user question, paste them into the prompt, and ask the "
        "LLM to answer using that context. Analogy: open-book exam—the book is your uploaded PDF."
    )
    doc.add_heading("Why RAG?", 2)
    doc.add_paragraph(
        "Without RAG, the model might hallucinate clauses or miss your actual terms. RAG grounds answers in "
        "retrieved text (with citations in this app)."
    )
    doc.add_heading("RAG flow in this project", 2)
    doc.add_paragraph(
        "1) Upload: document_service.extract_text → MongoDB documents.extracted_text → rag_service.index_document_chunks "
        "deletes old chunks for that document_id, chunks text, optionally embeds each chunk, inserts document_chunks."
    )
    doc.add_paragraph(
        "2) Chat: rag_service.retrieve loads all chunks for user_id (up to 3000), scores each chunk, top_k=5, "
        "build_context concatenates labeled chunks, ai_service.legal_chat sends prompt to Gemini or fallback."
    )
    add_mermaid_note(
        doc,
        "RAG pipeline (this codebase)",
        """
Upload
  -> save file (document_service.save_upload)
  -> extract_text (pypdf / docx / image placeholder)
  -> insert documents collection
  -> chunk_text (900/180 overlap)
  -> for each chunk: embed (Gemini) or None
  -> insert document_chunks (content, token_set, embedding)

Chat message
  -> tokenize query
  -> embed query (or None)
  -> for each chunk: score = 0.65*cosine + 0.35*lexical
  -> top 5 -> build_context -> legal_chat prompt
""",
    )
    doc.add_heading("Embeddings and similarity", 2)
    doc.add_paragraph(
        "Embeddings are fixed-length vectors representing meaning. rag_service._cosine computes cosine similarity "
        "between query embedding and chunk embedding. Lexical score is Jaccard-like: intersection of token sets / "
        "query token count. Chunk overlap (180 chars) reduces the chance that a sentence is split badly between chunks."
    )
    doc.add_paragraph(
        "Files: chunking and embedding in backend/app/services/rag_service.py; prompts in "
        "backend/app/services/ai_service.py (legal_chat). Retrieval orchestration in "
        "backend/app/modules/chat/routes.py (retrieve → build_context → legal_chat)."
    )

    # --- Section 6 API ---
    doc.add_heading("Section 6: API Documentation", 1)
    add_table(
        doc,
        ["Method", "Path", "Auth", "Body / Query", "Response / notes"],
        [
            ["GET", "/health", "No", "-", '{"status":"ok"}'],
            ["POST", "/auth/register", "No", "UserCreate JSON", "UserResponse; 400 if email exists"],
            ["POST", "/auth/login", "No", "UserLogin", "TokenResponse (access+refresh)"],
            ["POST", "/auth/refresh", "No", '{ "refresh_token" }', "TokenResponse; note key name must match backend"],
            ["GET", "/auth/me", "Bearer", "-", "UserResponse; 401 invalid token"],
            ["PATCH", "/users/me", "Bearer", "UserUpdate", "UserResponse"],
            ["POST", "/documents/upload", "Bearer", "multipart file", "DocumentResponse; 400 bad ext"],
            ["GET", "/documents", "Bearer", "-", "DocumentResponse[]"],
            ["GET", "/documents/{id}", "Bearer", "-", "DocumentResponse; 404"],
            ["DELETE", "/documents/{id}", "Bearer", "-", "Deletes clauses, risks, summaries, chunks; 404"],
            ["POST", "/documents/{id}/analyze", "Bearer", "-", "Runs AI heuristics; re-indexes RAG; notifications"],
            ["GET", "/documents/{id}/clauses", "Bearer", "-", "ClauseResponse[]"],
            ["GET", "/documents/{id}/risks", "Bearer", "-", "RiskResponse[]"],
            ["GET", "/documents/{id}/summary", "Bearer", "-", "SummaryResponse; 404 if no summary"],
            ["POST", "/documents/{id}/clauses/explain", "Bearer", "ClauseExplainRequest", "Heuristic JSON from ai_service"],
            ["POST", "/chat/sessions", "Bearer", "ChatCreate", "{id, title}"],
            ["GET", "/chat/sessions", "Bearer", "-", "list sessions"],
            ["GET", "/chat/sessions/{id}", "Bearer", "-", "session with messages"],
            ["POST", "/chat/sessions/{id}/messages", "Bearer", "ChatMessageRequest", "ChatMessageResponse + RAG"],
            ["GET", "/search?q=", "Bearer", "q min length 1", "documents, chats, clauses, risks hits"],
            ["GET", "/notifications", "Bearer", "-", "notification list"],
        ],
    )
    doc.add_paragraph(
        "Frontend note: api.ts posts /auth/refresh with body key refresh_token; backend expects the same in "
        "payload.get('refresh_token') — they align."
    )

    # --- Section 7 DB ---
    doc.add_heading("Section 7: Database Design", 1)
    doc.add_paragraph(
        "MongoDB is schemaless; the application enforces shape in code. Logical collections:"
    )
    add_table(
        doc,
        ["Collection", "Key fields", "Purpose"],
        [
            ["users", "_id, email, password_hash, name, created_at", "Auth and profile"],
            ["documents", "_id, user_id, file_name, storage_path, extracted_text, analysis_status, risk_score", "Uploaded files metadata + text"],
            ["document_chunks", "user_id, document_id, chunk_index, content, token_set, embedding, file_name", "RAG storage"],
            ["clauses", "document_id, clause_type, content, explanation, risk_level", "Heuristic clause extraction"],
            ["risks", "document_id, severity, issue, recommendation, ...", "Risk rules output"],
            ["summaries", "document_id, plain_summary, obligations, ...", "Heuristic summary"],
            ["chats", "user_id, title, messages[], updated_at", "Chat history"],
            ["notifications", "user_id, title, message, created_at", "User notifications"],
        ],
    )
    doc.add_paragraph(
        "Relationships: document_id in clauses/risks/summaries/chunks is a string of Mongo ObjectId. "
        "Indexes: see db/indexes.py (unique email, compound user+time, etc.)."
    )
    add_mermaid_note(
        doc,
        "ER-style (conceptual)",
        """
users 1 --- * documents
documents 1 --- * clauses
documents 1 --- * risks
documents 1 --- * summaries
documents 1 --- * document_chunks
users 1 --- * chats
users 1 --- * notifications
""",
    )

    # --- Section 8 Execution flow ---
    doc.add_heading("Section 8: Execution Flow — User uploads document", 1)
    steps = [
        "POST /documents/upload hits upload_document in documents/routes.py.",
        "Extension checked against ALLOWED; else 400.",
        "document_service.save_upload writes bytes to uploads/{uuid}{ext}.",
        "document_service.extract_text: PDF via pypdf pages; DOCX via paragraphs; image returns placeholder string.",
        "Mongo insert documents with analysis_status 'uploaded'.",
        "rag_service.index_document_chunks: delete_many chunks for new id; chunk_text; per chunk _embed_text; insert_many.",
        "Notification 'Upload successful' inserted.",
        "DocumentResponse returned to client.",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}")

    # --- Section 9 Concepts ---
    doc.add_heading("Section 9: Important Concepts Explained", 1)
    concepts = [
        "Decorators: @router.get registers a path; @app.on_event('startup') runs code when server starts.",
        "async/await: I/O operations (DB, HTTP) yield control so other requests progress.",
        "Middleware: CORS wraps responses to allow browser cross-origin (here allow_origins=['*']).",
        "Dependency injection: Depends(get_current_user_id) tells FastAPI to run that function and pass result.",
        "Pydantic models: UserCreate validates email format and password length before your handler runs.",
        "Service layer: AIService, RAGService, DocumentService hold business logic separate from HTTP.",
        "Repository pattern: not formal here; Motor calls are inline in routes (could refactor to repositories).",
    ]
    for c in concepts:
        doc.add_paragraph(c, style="List Bullet")

    # --- Section 10 Security ---
    doc.add_heading("Section 10: Security", 1)
    doc.add_paragraph(
        "Passwords: hashed with passlib; never store plaintext. JWT: access token type 'access' required for protected "
        "routes; refresh token type 'refresh'. API security: Bearer header. CORS is permissive (\"*\")—tighten in production "
        "to your frontend origin. Input validation: Pydantic on JSON bodies; upload extension whitelist. "
        "No rate limiting in code—add reverse proxy or middleware for production. "
        "Files stored on disk under uploads/ — protect directory permissions and backups."
    )

    # --- Section 11 Deployment ---
    doc.add_heading("Section 11: Deployment", 1)
    doc.add_paragraph("Local backend: cd backend; python -m venv .venv; pip install -r requirements.txt; copy .env from .env.example; uvicorn app.main:app --reload")
    doc.add_paragraph("Local frontend: cd frontend; npm install; set VITE_API_BASE_URL=http://localhost:8000; npm run dev")
    doc.add_paragraph("Docker: docker compose up --build — exposes mongo 27017, backend 8000, frontend 5173 mapped to nginx 80 in compose.")
    doc.add_paragraph("Production: set strong JWT_SECRET, restrict CORS, use TLS termination, secure MongoDB, add monitoring and log aggregation.")

    # --- Section 12 Errors ---
    doc.add_heading("Section 12: Error Handling", 1)
    doc.add_paragraph(
        "Gemini: _call_gemini retries once on 429/5xx; then returns None and fallbacks apply. "
        "HTTPException for 400/401/404 from routes. Motor/BSON errors can surface as 500 if not caught—extend with "
        "exception handlers for production. Logging: logging.getLogger in AIService for Gemini warnings."
    )

    # --- Section 13 Optimization ---
    doc.add_heading("Section 13: Optimization Opportunities", 1)
    doc.add_paragraph(
        "retrieve() loads up to 3000 chunks per user into memory—use vector index + limit by document or approximate NN. "
        "Embed API calls sequential per chunk on upload—batch if API allows. "
        "Clause extraction is keyword-based—replace with LLM structured output when budget allows. "
        "Image OCR placeholder—integrate Tesseract or Gemini vision. "
        "Add Redis for session rate limits and optional cache of embeddings."
    )

    # --- Section 14 Interview ---
    doc.add_heading("Section 14: Interview Preparation", 1)

    def iq(topic: str, qa: list[tuple[str, str]]) -> None:
        doc.add_paragraph(topic, style="Heading 2")
        for q, a in qa:
            doc.add_paragraph(f"Q: {q}", style="Heading 4")
            doc.add_paragraph(a)

    iq(
        "RAG",
        [
            (
                "What is RAG?",
                "Retrieval-Augmented Generation combines search over your documents with an LLM prompt so answers "
                "are grounded in retrieved chunks. In this project, chunks live in MongoDB and are ranked by hybrid score.",
            ),
            (
                "Fine-tuning vs RAG?",
                "Fine-tuning changes model weights on training data; RAG keeps the model frozen and injects context at "
                "query time. RAG is better for frequently changing private documents without retraining.",
            ),
            (
                "How does vector search work here?",
                "Not a dedicated ANN index—Mongo stores vectors; Python computes cosine vs query embedding for each chunk. "
                "Scales poorly but fine for MVP.",
            ),
            (
                "What are embeddings?",
                "Dense vectors from an embedding model so semantically similar texts have similar vectors; cosine measures angle similarity.",
            ),
            (
                "What is chunk overlap?",
                "Successive text windows share characters (180 here) so sentences at boundaries appear in two chunks, improving recall.",
            ),
        ],
    )
    iq(
        "FastAPI / Python async",
        [
            (
                "Why FastAPI for ML backends?",
                "Native async, validation, and OpenAPI; pairs well with async HTTP and DB drivers.",
            ),
            (
                "What does async def do?",
                "The function returns a coroutine; await yields during I/O so the server can handle other requests.",
            ),
        ],
    )
    iq(
        "MongoDB",
        [
            (
                "Why Mongo for this app?",
                "Flexible schema for evolving legal fields and embedded vectors in documents without migrations.",
            ),
            (
                "Indexes used?",
                "Unique email; compound user_id + created_at on documents; chunk indexes per user/document.",
            ),
        ],
    )
    iq(
        "Authentication",
        [
            (
                "How does JWT auth work here?",
                "Login returns signed tokens; client sends Authorization: Bearer access token; deps decode and check type 'access'.",
            ),
            (
                "Why refresh tokens?",
                "Short-lived access tokens limit exposure; refresh gets new tokens without re-entering password.",
            ),
        ],
    )
    iq(
        "Architecture (system design)",
        [
            (
                "How would you scale RAG here?",
                "Move embeddings to a vector DB with ANN; background workers for chunking; cache frequent queries; shard Mongo by tenant.",
            ),
            (
                "Single vs microservices?",
                "Current monolith is fine for MVP; split if teams need independent deploys or CPU-heavy ML isolation.",
            ),
        ],
    )
    iq(
        "Embeddings and search",
        [
            (
                "What is semantic vs lexical search?",
                "Lexical matches words (regex or token overlap). Semantic uses vectors so synonyms and paraphrases can match.",
            ),
            (
                "Why hybrid score 0.65/0.35?",
                "Weights vector vs keyword signals so retrieval still works if embeddings fail or are absent.",
            ),
        ],
    )

    # --- Section 15 Code walkthrough ---
    doc.add_heading("Section 15: Important Code Walkthrough", 1)
    doc.add_heading("15.1 backend/app/main.py (application entry)", 2)
    doc.add_paragraph(
        "Line 1–3: Import FastAPI, CORS, settings, lifecycle DB helpers, and every APIRouter. "
        "Line 15: app = FastAPI(title=settings.app_name) registers the ASGI app and sets OpenAPI title. "
        "Lines 17–23: CORSMiddleware allows browser clients from any origin (development-friendly; narrow in production). "
        "Lines 26–31: startup event awaits connect_to_mongo() then ensure_indexes() so collections have unique email and query-friendly indexes. "
        "Lines 34–36: shutdown closes the Motor client. "
        "Lines 39–41: /health for load balancers. "
        "Lines 44–50: include_router mounts auth (two routers), documents, analysis, chat, search, notifications."
    )
    doc.add_heading("15.2 backend/app/services/rag_service.py", 2)
    doc.add_paragraph(
        "chunk_text: normalizes whitespace to single spaces, slices [start:end] with end=start+chunk_size, advances start to end-overlap. "
        "_tokenize: regex [a-zA-Z0-9_]{3,} lowercased to a set for overlap scoring. "
        "_embed_text: POST to text-embedding-004 with truncated text; returns None on missing key or HTTP error. "
        "index_document_chunks: delete prior chunks for document_id; build docs with embedding + token_set; insert_many. "
        "retrieve: fetches all chunks for user (cap 3000 in code), computes lexical + cosine hybrid, sorts, returns top_k. "
        "build_context: prefixes each chunk with file name and index for traceability in prompts."
    )
    doc.add_heading("15.3 backend/app/services/ai_service.py", 2)
    doc.add_paragraph(
        "_call_gemini: builds REST URL with gemini_model and API key; POST JSON contents.parts.text; retry once on rate limit/server errors. "
        "extract_clauses: not LLM—loops CLAUSE_KEYWORDS and matches sentences containing keywords. "
        "detect_risks: regex checklist against full corpus; accumulates risks or emits low-severity default. "
        "summarize_contract: first sentences / keyword filters—heuristic MVP. "
        "legal_chat: if context exists, prepends up to 4000 chars; else general-knowledge prompt; fallback strings if Gemini absent."
    )
    doc.add_heading("15.4 backend/app/modules/chat/routes.py (send_message)", 2)
    doc.add_paragraph(
        "Loads chat session for user; rag_service.retrieve(..., top_k=5); build_context; ai_service.legal_chat; "
        "if chunks exist, citations from rag_service.build_citations; appends user + assistant messages with ISO timestamps; "
        "updates chats collection with full messages array."
    )
    doc.add_heading("15.5 frontend/src/lib/api.ts", 2)
    doc.add_paragraph(
        "Axios baseURL from VITE_API_BASE_URL or localhost:8000. Request interceptor adds Bearer accessToken from localStorage. "
        "refreshAccessToken POSTs /auth/refresh with refresh_token body key (must match backend)."
    )

    add_code_block(
        doc,
        """# Hybrid retrieval (conceptual — see rag_service.py)
lexical = len(query_tokens & chunk_tokens) / len(query_tokens)
vector = cosine(query_embedding, chunk_embedding)  # 0 if either missing
score = 0.65 * vector + 0.35 * lexical
""",
    )

    doc.add_heading("Appendix A: Source files inventory (project-authored)", 2)
    doc.add_paragraph(
        "The following lists first-party source files under backend/app and frontend/src (excluding node_modules, .venv). "
        "Config/tooling: docker-compose.yml, backend/Dockerfile, backend/requirements.txt, backend/.env.example, "
        "frontend/package.json, frontend/vite.config.ts, frontend/index.html (if present)."
    )
    add_table(
        doc,
        ["Path", "Role"],
        [
            ["backend/app/__init__.py", "Package marker"],
            ["backend/app/main.py", "FastAPI app assembly"],
            ["backend/app/core/config.py", "Environment settings"],
            ["backend/app/core/deps.py", "Auth + DB dependencies"],
            ["backend/app/core/security.py", "JWT + password hashing"],
            ["backend/app/db/mongo.py", "Motor connection"],
            ["backend/app/db/indexes.py", "Index creation"],
            ["backend/app/schemas/common.py", "Auth/user Pydantic models"],
            ["backend/app/schemas/legal.py", "Legal/chat Pydantic models"],
            ["backend/app/modules/auth/routes.py", "Auth HTTP routes"],
            ["backend/app/modules/documents/routes.py", "Upload, CRUD, analyze"],
            ["backend/app/modules/analysis/routes.py", "Clauses, risks, summary, explain"],
            ["backend/app/modules/chat/routes.py", "Chat sessions and messages"],
            ["backend/app/modules/search/routes.py", "Global search"],
            ["backend/app/modules/notifications/routes.py", "Notifications list"],
            ["backend/app/services/ai_service.py", "Gemini + heuristics"],
            ["backend/app/services/rag_service.py", "RAG chunk/embed/retrieve"],
            ["backend/app/services/document_service.py", "File save + text extraction"],
            ["frontend/src/main.tsx", "React root mount"],
            ["frontend/src/App.tsx", "Routes + AuthProvider"],
            ["frontend/src/lib/api.ts", "Axios client"],
            ["frontend/src/lib/utils.ts", "cn() helper for Tailwind"],
            ["frontend/src/lib/types.ts", "Shared TS types"],
            ["frontend/src/hooks/use-auth.tsx", "Auth context"],
            ["frontend/src/pages/auth-pages.tsx", "Login/register UI"],
            ["frontend/src/pages/landing-page.tsx", "Marketing landing"],
            ["frontend/src/pages/workspace-pages.tsx", "Dashboard, chat, upload, documents, analysis, profile, settings"],
            ["frontend/src/components/app-sidebar.tsx", "Navigation sidebar"],
            ["frontend/src/components/chat-bubble.tsx", "Chat UI"],
            ["frontend/src/components/upload-zone.tsx", "Upload UI"],
            ["frontend/src/components/risk-badge.tsx", "Risk display"],
            ["frontend/src/components/feature-card.tsx", "Landing components"],
            ["frontend/src/components/marketing-navbar.tsx", "Landing nav"],
            ["frontend/src/components/ui/*", "Button, card, input, skeleton primitives"],
        ],
    )

    # --- Section 16 Glossary ---
    doc.add_heading("Section 16: Glossary", 1)
    gloss = [
        ("Embedding", "Vector representation of text for similarity."),
        ("Vector database", "Specialized store for ANN search; here MongoDB holds arrays without a separate ANN engine."),
        ("Chunking", "Splitting long text into smaller pieces for retrieval."),
        ("Chunk overlap", "Shared characters between adjacent chunks to avoid cutting sentences in half."),
        ("Retriever", "Code path that selects relevant chunks for a query (RAGService.retrieve)."),
        ("Prompt engineering", "Crafting instructions and context for the LLM (legal_chat prompts)."),
        ("LLM", "Large language model; here Google Gemini via generateContent REST API."),
        ("Inference", "Running the model to produce output from input."),
        ("Token", "LLM text unit; here token_set also means regex-extracted word-like strings for lexical scoring."),
        ("Context window", "Maximum tokens the model can attend to; legal_chat sends ~4000 chars of context."),
        ("Cosine similarity", "Dot product divided by vector lengths; 1 = same direction, 0 = unrelated."),
        ("ASGI", "Asynchronous Server Gateway Interface; Uvicorn runs FastAPI ASGI app."),
        ("Motor", "Async Python driver for MongoDB used with async route handlers."),
        ("Pydantic", "Data validation via Python types; defines request/response bodies."),
        ("JWT", "JSON Web Token; signed JSON claims for stateless auth."),
        ("OCR", "Optical character recognition; image flow is a placeholder in document_service.py."),
        ("CORS", "Cross-Origin Resource Sharing; browser security for API calls from other origins."),
        ("Environment variables", "Secrets and config (e.g. .env) loaded by pydantic-settings in config.py."),
    ]
    for term, expl in gloss:
        doc.add_paragraph(f"{term}: {expl}")

    doc.add_heading("Appendix B: Configuration reference", 1)
    add_table(
        doc,
        ["Variable", "Purpose"],
        [
            ["MONGO_URI", "Mongo connection string"],
            ["MONGO_DB_NAME", "Database name (default legal_ai)"],
            ["JWT_SECRET", "Signing key for JWT — must be long and secret in prod"],
            ["ACCESS_TOKEN / REFRESH_TOKEN exp", "Configured in settings; login uses fixed minutes in code paths — verify consistency when changing"],
            ["UPLOAD_DIR", "Directory for saved uploads"],
            ["GEMINI_API_KEY", "Enables Gemini chat and embeddings"],
            ["GEMINI_MODEL", "e.g. gemini-1.5-flash for generateContent"],
        ],
    )

    doc.add_heading("Appendix C: Implementation notes and known quirks", 1)
    doc.add_paragraph(
        "Password hashing uses passlib with pbkdf2_sha256 (see security.py), not bcrypt, despite passlib[bcrypt] in requirements. "
        "Token expiry: auth/routes.py uses literal 30 and 60*24*7 minutes in create_token calls; align with config settings if you change env defaults. "
        "LangChain, Redis, PostgreSQL, and dedicated vector DBs are not used—document accordingly when describing the stack. "
        "RAG without GEMINI_API_KEY: embeddings are None, so vector part of the score is zero; lexical matching still runs."
    )

    doc.add_paragraph(
        "End of document. For authoritative behavior, read the cited Python and TypeScript files in the repository."
    )

    doc.save(out_path)
    print(f"Wrote: {out_path}")


if __name__ == "__main__":
    main()
